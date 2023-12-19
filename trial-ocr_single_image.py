#!/usr/bin/env python3
import cv2
from os import environ, path, remove
import pytesseract as pts
from sys import argv
import docx
from fpdf import FPDF


# OUTPUT MODE ['print', 'file'] # output directory for 'file' mode
output_mode = 'file'
output_directory =  '.'


# INPUT FILE test pdf file location
image_test_directory = 'test_files/images/'
test_file = 'kidanewolde.png'


# TRAINING DATA directory (if not provieded uses tesseracts default dir)
# TESSDATA_PREFIX='/home/menelikberhan/amharic_ocr/test_data' (doesnt work)
environ['TESSDATA_PREFIX'] = '/home/menelikberhan/amharic_ocr/training_data'

################ END OF CONFIG #################




################ PROGRAM ENTRY #################

output_directory = path.abspath(output_directory)

image_file_path = image_test_directory + test_file

# check if input file exists
if not path.exists(image_file_path):
    print("Error: Specified input file '{}' doesn't exist".format(image_file_path))
    exit(1)

# evaluate absolute file path and remove file if it exists for 'file' output mode
if output_mode == 'file':
    output_file = test_file.split('.')[0] + '_output.txt'
    output_file_path = output_directory + '/' + output_file

    if not path.exists(output_directory):
        print("Error: Specified output directory '{}' doesn't exist".format(output_directory))
        exit(1)

    if path.exists(output_file_path):
        print("Warning: Output file with same name exists.")
        remove(output_file_path)
        print("Removed previous output file '{}' ".format(output_file_path))

# load image
image = cv2.imread(image_file_path)


# cv2 uses BGR, so convert to RGB
rgb_image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)


# config tesseract options
lang = 'amh'
psm = 1         # image processing option (1 & 3 for full page). refer doc.
options = "-l {} --psm {}".format(lang, psm)


# other formats image_to_[...] - 'data' with dict option, pdf, box ...
text = pts.image_to_string(rgb_image, config=options)

# out put based on output_mode
if output_mode == 'print':
    print(text)
elif output_mode == 'file':
    # with open(output_file_path, 'w', encoding='utf-8') as file:
    #     file.write(text)

    doc = docx.Document()
    par = doc.add_paragraph().add_run(text)
    par.font.name = 'Abyssinica SIL'
    doc.save('test.docx')
