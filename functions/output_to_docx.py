#!/usr/bin/env python3
import docx
from . import write_dict


def write_to_docx(text, output_file_path, document=None, **args):
    """Writes OCR'ed text to MS word document.
    
    Args:
        text (str): output from OCR by tesseract.
        output_file_path (str): path (including file name) of output file.
        document (Document): a Document object loaded from docx to add pages to.
            If not provided a new one will be created.
        **args (dict): dictionary of parameters (font, layout ...)
            for formatting document.

    Retruns:
        Document: a Document object with newly added page containing text.

    """
    # TODO - check if saving document after treshold no. of pages speeds up process
    document = document if document else docx.Document()
    par = document.add_paragraph().add_run(text)
    par.font.name = args.get('font_name', write_dict.get('font_name_def'))

    if args.get('save'):
        document.save(output_file_path)

    # TODO - handle other args for MSword formatiing

    return (document)
